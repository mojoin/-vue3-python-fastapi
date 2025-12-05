# -*- coding: utf-8 -*-
"""
地表沉降监测系统后端服务
功能：提供数据解析、变形指标计算（沉降、水平位移、坡度、曲率）、以及自动化监测报告生成。
作者：AI助手 (人工精细化注释版)
"""

import uvicorn
from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from typing import List, Optional, Dict, Any
import numpy as np
import pandas as pd
import io
import os
import re
import math
from datetime import datetime

# 引入报告生成相关的库
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib

# 设置 matplotlib 后端为 'Agg'
matplotlib.use('Agg')
plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial Unicode MS', 'sans-serif']
plt.rcParams['axes.unicode_minus'] = False

app = FastAPI(title="地表沉降监测分析系统")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- 业务逻辑辅助函数 ---

def is_date_like(s: Any) -> bool:
    s = str(s).strip()
    if not s or s.lower() == 'nan': return False
    return bool(re.search(r'\d{2,4}[/-]\d{1,2}[/-]\d{1,2}', s))

def clean_float(val: Any) -> float:
    try:
        return float(val)
    except:
        return 0.0

def generate_report_docx(analysis_data: Dict):
    """
    报告生成器 (精简版，包含核心图表)
    """
    doc = Document()
    
    # 标题
    heading = doc.add_heading('地表沉降监测分析报告', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"报告生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"数据源类型: {'全站仪 (XYZ三维数据)' if analysis_data['data_type'] == 'total_station' else '水准仪 (Z一维沉降数据)'}")
    
    # 1. 监测概况
    doc.add_heading('1. 监测概况', level=1)
    p = doc.add_paragraph()
    p.add_run(f"本次监测共涉及 {len(analysis_data['point_ids'])} 个测点，")
    p.add_run(f"监测时间跨度为 {analysis_data['dates'][0]} 至 {analysis_data['dates'][-1]}。")

    # 2. 变形数据分析
    doc.add_heading('2. 变形数据分析', level=1)
    doc.add_heading('2.1 关键变形指标统计', level=2)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '指标项'
    hdr[1].text = '极值数值'
    hdr[2].text = '发生测点'
    hdr[3].text = '发生日期'
    
    # 填充表格
    metrics = analysis_data['metrics']
    data = [
        ('最大累计下沉', f"{metrics['max_subsidence']['val']:.2f} mm", metrics['max_subsidence']['point'], metrics['max_subsidence']['date']),
        ('最大局部坡度', f"{metrics['max_slope']['val']:.4f}", metrics['max_slope']['point'], "区间计算")
    ]
    
    if analysis_data['data_type'] == 'total_station':
        data.insert(1, ('最大水平位移', f"{metrics['max_horizontal']['val']:.2f} mm", metrics['max_horizontal']['point'], metrics['max_horizontal']['date']))
        
    for item in data:
        row = table.add_row().cells
        row[0].text = item[0]
        row[1].text = item[1]
        row[2].text = str(item[2])
        row[3].text = str(item[3])

    # 2.2 趋势图
    doc.add_heading('2.2 变形趋势图', level=2)
    
    # 绘图辅助函数
    def add_plot_to_doc(series_key, title, y_label):
        plt.figure(figsize=(9, 4))
        last_date = analysis_data['dates'][-1]
        series_data = analysis_data[series_key]['series']
        last_series = next((s for s in series_data if s['name'] == last_date), None)
        
        if last_series:
            x_points = analysis_data['point_ids']
            y_values = last_series['data']
            plt.plot(x_points, y_values, marker='o', markersize=4, linestyle='-', label=last_date)
            plt.title(f"{title} ({last_date})")
            plt.xlabel('测点编号')
            plt.ylabel(y_label)
            plt.grid(True, linestyle='--', alpha=0.6)
            plt.legend()
            plt.tight_layout()
            
            img_buffer = io.BytesIO()
            plt.savefig(img_buffer, format='png')
            img_buffer.seek(0)
            doc.add_picture(img_buffer, width=Inches(6))
            doc.add_paragraph(f'图: {title}').alignment = WD_ALIGN_PARAGRAPH.CENTER
            plt.close()

    # 插入沉降图
    add_plot_to_doc('chart_legend_date_z', '各测点累计沉降分布', '累计沉降 (mm)')
    
    # 如果是全站仪，插入水平位移图
    if analysis_data['data_type'] == 'total_station':
        doc.add_paragraph("注：水平位移为XY合成矢量模长。")
        add_plot_to_doc('chart_legend_date_h', '各测点累计水平位移分布', '累计水平位移 (mm)')

    # 3. 结论
    doc.add_heading('3. 结论与建议', level=1)
    max_sub = abs(metrics['max_subsidence']['val'])
    conclusion = "本期监测数据表明，"
    if max_sub < 10: conclusion += "场地变形较小，处于稳定状态。"
    elif max_sub < 30: conclusion += "场地出现一定程度沉降，需持续关注。"
    else: conclusion += "局部区域沉降显著，建议加密监测。"
    doc.add_paragraph(conclusion)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# 获取端口配置
port = int(os.environ.get("PORT", 8001))
host = "0.0.0.0"

def parse_monitor_csv_v3(content: bytes):
    """
    CSV 解析核心引擎 (V3版 - 增强版)。
    """
    
    # 1. 尝试解码
    df_raw = None
    for encoding in ['utf-8', 'gbk', 'gb18030']:
        try:
            df_raw = pd.read_csv(io.BytesIO(content), header=None, encoding=encoding)
            break
        except:
            continue
    if df_raw is None:
        raise ValueError("无法识别文件编码")

    # 2. 定位表头
    header_row_idx = -1
    is_total_station = False
    
    for i in range(min(20, len(df_raw))):
        row_vals = [str(x).upper().strip() for x in df_raw.iloc[i].tolist()]
        cnt_x = row_vals.count('X')
        cnt_y = row_vals.count('Y')
        cnt_z = row_vals.count('Z')
        
        if cnt_x >= 2 and cnt_y >= 2 and cnt_z >= 2:
            header_row_idx = i
            is_total_station = True
            break
        if cnt_z >= 2:
            header_row_idx = i
            is_total_station = False
            
    if header_row_idx == -1: header_row_idx = 2

    # 3. 解析映射
    header_row = [str(x).upper().strip() for x in df_raw.iloc[header_row_idx].tolist()]
    date_row_idx = header_row_idx - 1
    date_row = [str(x).strip() for x in df_raw.iloc[date_row_idx].tolist()] if date_row_idx >= 0 else []

    dates = []
    point_ids = []
    raw_data_map = {} 
    col_mappings = [] 
    
    if is_total_station:
        for i in range(len(header_row) - 2):
            if header_row[i] == 'X' and header_row[i+1] == 'Y' and header_row[i+2] == 'Z':
                d_str = date_row[i] if i < len(date_row) else f"Date_{i}"
                if is_date_like(d_str):
                    col_mappings.append({'date': d_str, 'x': i, 'y': i+1, 'z': i+2})
                    if d_str not in dates: dates.append(d_str)
    else:
        z_indices = [i for i, v in enumerate(header_row) if v == 'Z']
        d_candidates = [d for d in date_row if is_date_like(d)]
        count = min(len(z_indices), len(d_candidates))
        for i in range(count):
            d_str = d_candidates[i]
            col_mappings.append({'date': d_str, 'x': -1, 'y': -1, 'z': z_indices[i]})
            if d_str not in dates: dates.append(d_str)

    # 读取数据
    data_start_idx = header_row_idx + 1
    for _, row in df_raw.iloc[data_start_idx:].iterrows():
        pid = str(row[0]).strip()
        if not pid or pid.lower() == 'nan': continue
        point_ids.append(pid)
        
        for mapping in col_mappings:
            d = mapping['date']
            if d not in raw_data_map: raw_data_map[d] = {}
            z_val = clean_float(row[mapping['z']])
            x_val = clean_float(row[mapping['x']]) if mapping['x'] != -1 else 0.0
            y_val = clean_float(row[mapping['y']]) if mapping['y'] != -1 else 0.0
            raw_data_map[d][pid] = {"x": x_val, "y": y_val, "z": z_val}

    # 4. 横坐标计算：每个点到第一个点的距离（直线长度）
    distances_to_first = []
    
    if point_ids:
        # 获取起点的参考坐标（使用第一个日期的数据作为基准）
        ref_date = dates[0]
        p0_data = raw_data_map.get(ref_date, {}).get(point_ids[0])
        
        # 即使第一个点在第一期没有数据，也要尝试寻找它的位置或者初始化为0
        p0_x, p0_y = 0.0, 0.0
        has_p0_coords = False
        
        if is_total_station and p0_data:
            p0_x, p0_y = p0_data["x"], p0_data["y"]
            has_p0_coords = True

        for i, pid in enumerate(point_ids):
            if i == 0:
                distances_to_first.append(0.0) # 第一个点距离自己为0
                continue
            
            calculated_dist = None
            
            # 如果是全站仪数据，尝试通过坐标计算距离
            if is_total_station and has_p0_coords:
                # 优先使用第一期数据计算
                p_curr = raw_data_map.get(ref_date, {}).get(pid)
                
                # 如果第一期该点缺失，尝试遍历所有日期寻找该点坐标来估算位置
                if not p_curr:
                    for d in dates:
                        if pid in raw_data_map.get(d, {}):
                            p_curr = raw_data_map[d][pid]
                            break
                
                if p_curr:
                    dx = p_curr["x"] - p0_x
                    dy = p_curr["y"] - p0_y
                    # 欧氏距离，结果必为正数
                    calculated_dist = math.sqrt(dx*dx + dy*dy)
            
            # 如果无法计算（水准仪数据 或 坐标缺失），使用默认逻辑
            # 这里沿用默认逻辑：距离 = 上一个点的距离 + 10m
            if calculated_dist is None:
                prev_dist = distances_to_first[-1]
                calculated_dist = prev_dist + 10.0
                
            distances_to_first.append(calculated_dist)
    
    # 5. 纵坐标数据准备：提取原始数据
    raw_values_by_date = {}
    for d in dates:
        current_points_data = raw_data_map.get(d, {})
        raw_values_by_date[d] = {
            "z": [],
            "x": [],
            "y": []
        }
        
        for pid in point_ids:
            curr = current_points_data.get(pid, {"x":0,"y":0,"z":0})
            raw_values_by_date[d]["z"].append(curr["z"])
            raw_values_by_date[d]["x"].append(curr["x"])
            raw_values_by_date[d]["y"].append(curr["y"])
    
    # 6. 计算累计变形量（纵坐标核心逻辑）
    # 逻辑：当前值 - 初始值 (或 初始值 - 当前值，视具体指标定义)
    cumulative_diff_by_date = {}
    
    initial_date = dates[0]
    initial_vals = raw_values_by_date[initial_date]
    
    for d in dates:
        cumulative_diff_by_date[d] = {
            "z": [], "h": [], "x": [], "y": []
        }
        
        curr_vals = raw_values_by_date[d]
        
        for j in range(len(point_ids)):
            # Z 沉降：通常定义为 初始高程 - 当前高程 (下沉为正)
            z_diff = initial_vals["z"][j] - curr_vals["z"][j]
            cumulative_diff_by_date[d]["z"].append(z_diff)
            
            # 水平位移相关
            if is_total_station:
                # X方向累计位移: 当前 - 初始
                x_diff = curr_vals["x"][j] - initial_vals["x"][j]
                cumulative_diff_by_date[d]["x"].append(x_diff)
                
                # Y方向累计位移: 当前 - 初始
                y_diff = curr_vals["y"][j] - initial_vals["y"][j]
                cumulative_diff_by_date[d]["y"].append(y_diff)
                
                # 平面合位移: sqrt(dx^2 + dy^2)
                h_diff = math.sqrt(x_diff**2 + y_diff**2)
                cumulative_diff_by_date[d]["h"].append(h_diff)
            else:
                cumulative_diff_by_date[d]["x"].append(0.0)
                cumulative_diff_by_date[d]["y"].append(0.0)
                cumulative_diff_by_date[d]["h"].append(0.0)

    # 7. 生成图表序列数据
    series_z_time_legend = []
    series_h_time_legend = []
    series_x_time_legend = []
    series_y_time_legend = []
    
    for d in dates:
        series_z_time_legend.append({"name": d, "type": "line", "data": cumulative_diff_by_date[d]["z"]})
        series_h_time_legend.append({"name": d, "type": "line", "data": cumulative_diff_by_date[d]["h"]})
        series_x_time_legend.append({"name": d, "type": "line", "data": cumulative_diff_by_date[d]["x"]})
        series_y_time_legend.append({"name": d, "type": "line", "data": cumulative_diff_by_date[d]["y"]})
    
    # 8. 计算统计指标（寻找最大值）
    max_sub = {"val": 0, "point": "--", "date": "--"}
    max_hor = {"val": 0, "point": "--", "date": "--"}
    max_slp = {"val": 0, "point": "--", "date": "--"}
    max_crv = {"val": 0, "point": "--", "date": "--"}
    
    for d in dates:
        # 使用上面计算好的累计值来找最大值
        z_data = cumulative_diff_by_date[d]["z"]
        h_data = cumulative_diff_by_date[d]["h"]
        
        for j, pid in enumerate(point_ids):
            # 最大沉降
            if z_data[j] > max_sub["val"]: # 假设沉降为正
                max_sub["val"] = z_data[j]
                max_sub["point"] = pid
                max_sub["date"] = d
            
            # 最大水平位移
            if h_data[j] > max_hor["val"]:
                max_hor["val"] = h_data[j]
                max_hor["point"] = pid
                max_hor["date"] = d

    # 9. 计算坡度和曲率（基于累计沉降曲线的形态）
    for d in dates:
        current_z_cumul = cumulative_diff_by_date[d]["z"]
        current_slopes = []
        
        for i in range(len(point_ids) - 1):
            z1, z2 = current_z_cumul[i], current_z_cumul[i+1]
            
            # 距离差：横坐标的差值
            dist = distances_to_first[i+1] - distances_to_first[i] if i+1 < len(distances_to_first) else 10.0
            if dist <= 0: dist = 10.0
            
            # 坡度 = 沉降差 / 距离
            slope = (z2 - z1) / dist
            current_slopes.append(slope)
            if abs(slope) > abs(max_slp["val"]):
                max_slp["val"] = slope
                max_slp["point"] = f"{point_ids[i]}-{point_ids[i+1]}"
                max_slp["date"] = d
        
        # 曲率 = 坡度差
        for i in range(len(current_slopes) - 1):
            k = current_slopes[i+1] - current_slopes[i]
            if abs(k) > abs(max_crv["val"]):
                max_crv["val"] = k
                max_crv["point"] = f"{point_ids[i+1]}附近"
                max_crv["date"] = d

    return {
        "data_type": "total_station" if is_total_station else "leveling",
        "point_ids": point_ids,
        "dates": dates,
        "metrics": {
            "max_subsidence": max_sub,
            "max_horizontal": max_hor,
            "max_slope": max_slp,
            "max_curvature": max_crv
        },
        "chart_time_legend_distance_z": { "xAxis": distances_to_first, "series": series_z_time_legend, "xAxisName": "与第一个点的距离 (m)" },
        "chart_time_legend_distance_h": { "xAxis": distances_to_first, "series": series_h_time_legend, "xAxisName": "与第一个点的距离 (m)" },
        "chart_time_legend_distance_x": { "xAxis": distances_to_first, "series": series_x_time_legend, "xAxisName": "与第一个点的距离 (m)" },
        "chart_time_legend_distance_y": { "xAxis": distances_to_first, "series": series_y_time_legend, "xAxisName": "与第一个点的距离 (m)" }
    }

# --- API 路由 ---

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    try:
        content = await file.read()
        result = parse_monitor_csv_v3(content)
        return {"code": 200, "data": result}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

@app.post("/api/report")
async def create_report(file: UploadFile = File(...)):
    try:
        content = await file.read()
        analysis_data = parse_monitor_csv_v3(content)
        report_stream = generate_report_docx(analysis_data)
        filename = f"Monitor_Report_{datetime.now().strftime('%Y%m%d')}.docx"
        return StreamingResponse(
            report_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"报告生成失败: {str(e)}")

@app.get("/files/{filename}")
async def download_file(filename: str):
    if filename not in ["全站仪.csv", "水准.csv"]:
        raise HTTPException(status_code=404, detail="文件不存在")
    if os.path.exists(filename):
        return FileResponse(filename)
    return HTTPException(status_code=404)

@app.get("/")
async def read_root():
    return FileResponse("./index.html")

@app.post("/api/export-academic-chart")
async def export_academic_chart(file: UploadFile = File(...), chart_type: str = Form(...)):
    try:
        content = await file.read()
        analysis_data = parse_monitor_csv_v3(content)
        
        # 设置Matplotlib学术风格
        plt.style.use('seaborn-v0_8-paper')
        plt.rcParams.update({
            'font.size': 12,
            'axes.titlesize': 14,
            'axes.labelsize': 12,
            'xtick.labelsize': 10,
            'ytick.labelsize': 10,
            'legend.fontsize': 10,
            'figure.figsize': (10, 6),
            'figure.dpi': 300,
            'lines.linewidth': 1.5,
            'lines.markersize': 5,
            'axes.grid': True,
            'grid.linestyle': '--',
            'grid.alpha': 0.7,
            'font.family': 'SimHei'
        })
        
        # 根据图表类型选择数据
        chart_config = {
            'z': {
                'series_key': 'chart_time_legend_distance_z',
                'title': '不同时间累计沉降分布',
                'y_label': '累计沉降 (mm)',
                'color_map': 'viridis'
            },
            'h': {
                'series_key': 'chart_time_legend_distance_h',
                'title': '不同时间累计水平位移分布',
                'y_label': '累计水平位移 (mm)',
                'color_map': 'plasma'
            },
            'x': {
                'series_key': 'chart_time_legend_distance_x',
                'title': '不同时间累计X方向位移分布',
                'y_label': '累计X方向位移 (mm)',
                'color_map': 'inferno'
            },
            'y': {
                'series_key': 'chart_time_legend_distance_y',
                'title': '不同时间累计Y方向位移分布',
                'y_label': '累计Y方向位移 (mm)',
                'color_map': 'magma'
            }
        }
        
        config = chart_config.get(chart_type, chart_config['z'])
        dataset = analysis_data[config['series_key']]
        
        # 创建图表
        fig, ax = plt.subplots(figsize=(10, 6))
        
        # 使用colormap为不同时间分配颜色
        colors = plt.cm.get_cmap(config['color_map'], len(dataset['series']))
        
        # 绘制每条曲线
        for i, series in enumerate(dataset['series']):
            ax.plot(dataset['xAxis'], series['data'], 
                   label=series['name'], 
                   color=colors(i),
                   marker='o',
                   linestyle='-')
        
        # 设置图表属性
        ax.set_title(config['title'], fontweight='bold')
        ax.set_xlabel(dataset.get('xAxisName', '与第一个点的距离 (m)'), fontweight='bold')
        ax.set_ylabel(config['y_label'], fontweight='bold')
        
        # 添加图例
        ax.legend(loc='center left', bbox_to_anchor=(1, 0.5), frameon=True, ncol=1)
        
        # 调整布局
        plt.tight_layout()
        
        # 保存到内存
        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', dpi=300, bbox_inches='tight')
        img_buffer.seek(0)
        plt.close()
        
        # 返回图片
        filename = f"academic_chart_{chart_type}_{datetime.now().strftime('%Y%m%d%H%M%S')}.png"
        return StreamingResponse(
            img_buffer,
            media_type="image/png",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"图表生成失败: {str(e)}")

if __name__ == "__main__":
    uvicorn.run(app, host=host, port=port)